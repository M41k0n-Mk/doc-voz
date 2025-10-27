"""
Testes unitários para o módulo de parsers.
"""
import pytest
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch
from docx import Document

from src.parsers.base import Parser
from src.parsers.txt_parser import TxtParser
from src.parsers.docx_parser import DocxParser
from src.parsers.factory import ParserFactory


class TestTxtParser:
    """Testes para TxtParser."""
    
    def test_parse_simple_text(self):
        """Testa parsing de texto simples."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Texto de teste\nSegunda linha")
            temp_path = Path(f.name)
        
        try:
            parser = TxtParser()
            result = parser.parse(temp_path)
            assert result == "Texto de teste\nSegunda linha"
        finally:
            temp_path.unlink()
    
    def test_parse_empty_file(self):
        """Testa parsing de arquivo vazio."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("")
            temp_path = Path(f.name)
        
        try:
            parser = TxtParser()
            result = parser.parse(temp_path)
            assert result == ""
        finally:
            temp_path.unlink()
    
    def test_parse_utf8_encoding(self):
        """Testa parsing com caracteres UTF-8."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("Texto com acentos: ção, não, coração")
            temp_path = Path(f.name)
        
        try:
            parser = TxtParser()
            result = parser.parse(temp_path)
            assert "ção" in result
            assert "não" in result
            assert "coração" in result
        finally:
            temp_path.unlink()
    
    def test_parse_nonexistent_file(self):
        """Testa parsing de arquivo inexistente."""
        parser = TxtParser()
        nonexistent_path = Path("/path/that/does/not/exist.txt")
        
        with pytest.raises(FileNotFoundError):
            parser.parse(nonexistent_path)


class TestDocxParser:
    """Testes para DocxParser."""
    
    @patch('src.parsers.docx_parser.Document')
    def test_parse_simple_docx(self, mock_document):
        """Testa parsing de documento simples."""
        # Mock do documento
        mock_doc = Mock()
        mock_paragraph1 = Mock()
        mock_paragraph1.text = "Primeiro parágrafo"
        mock_paragraph2 = Mock()
        mock_paragraph2.text = "Segundo parágrafo"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]
        mock_document.return_value = mock_doc
        
        parser = DocxParser()
        result = parser.parse(Path("fake.docx"))
        
        assert result == "Primeiro parágrafo\nSegundo parágrafo"
        mock_document.assert_called_once_with(Path("fake.docx"))
    
    @patch('src.parsers.docx_parser.Document')
    def test_parse_empty_paragraphs(self, mock_document):
        """Testa parsing ignorando parágrafos vazios."""
        mock_doc = Mock()
        mock_paragraph1 = Mock()
        mock_paragraph1.text = "Texto válido"
        mock_paragraph2 = Mock()
        mock_paragraph2.text = ""  # Parágrafo vazio
        mock_paragraph3 = Mock()
        mock_paragraph3.text = "Outro texto"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2, mock_paragraph3]
        mock_document.return_value = mock_doc
        
        parser = DocxParser()
        result = parser.parse(Path("fake.docx"))
        
        assert result == "Texto válido\nOutro texto"
    
    @patch('src.parsers.docx_parser.Document')
    def test_parse_only_empty_paragraphs(self, mock_document):
        """Testa parsing de documento só com parágrafos vazios."""
        mock_doc = Mock()
        mock_paragraph1 = Mock()
        mock_paragraph1.text = ""
        mock_paragraph2 = Mock()
        mock_paragraph2.text = "   "  # Só espaços
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]
        mock_document.return_value = mock_doc
        
        parser = DocxParser()
        result = parser.parse(Path("fake.docx"))
        
        # O parser retorna o que há nos parágrafos, incluindo espaços
        assert result == "   "


class TestParserFactory:
    """Testes para ParserFactory."""
    
    def test_get_parser_for_txt(self):
        """Testa criação de parser para .txt."""
        parser = ParserFactory.get_parser_for_path(Path("test.txt"))
        assert isinstance(parser, TxtParser)
    
    def test_get_parser_for_docx(self):
        """Testa criação de parser para .docx."""
        parser = ParserFactory.get_parser_for_path(Path("test.docx"))
        assert isinstance(parser, DocxParser)
    
    def test_get_parser_case_insensitive(self):
        """Testa criação de parser insensível a maiúsculas/minúsculas."""
        parser_txt = ParserFactory.get_parser_for_path(Path("test.TXT"))
        parser_docx = ParserFactory.get_parser_for_path(Path("test.DOCX"))
        
        assert isinstance(parser_txt, TxtParser)
        assert isinstance(parser_docx, DocxParser)
    
    def test_get_parser_unsupported_extension(self):
        """Testa erro para extensão não suportada."""
        with pytest.raises(ValueError, match="Extensão não suportada"):
            ParserFactory.get_parser_for_path(Path("test.pdf"))
    
    def test_get_parser_no_extension(self):
        """Testa erro para arquivo sem extensão."""
        with pytest.raises(ValueError, match="Extensão não suportada"):
            ParserFactory.get_parser_for_path(Path("test"))


class TestParserBase:
    """Testes para a classe base Parser."""
    
    def test_parser_is_abstract(self):
        """Testa que Parser não pode ser instanciada diretamente."""
        with pytest.raises(TypeError):
            Parser()
    
    def test_parser_subclass_must_implement_parse(self):
        """Testa que subclasses devem implementar parse."""
        class IncompleteParser(Parser):
            pass
        
        with pytest.raises(TypeError):
            IncompleteParser()